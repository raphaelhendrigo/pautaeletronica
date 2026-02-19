from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass
from datetime import datetime, timedelta, date
from pathlib import Path
from typing import Optional, Tuple, List

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.opc.exceptions import PackageNotFoundError


# =========================
# Utilidades bÃ¡sicas
# =========================

def _ws(s) -> str:
    if s is None or (isinstance(s, float) and pd.isna(s)):
        return ""
    return re.sub(r"\s+", " ", str(s)).strip()


def _strip_accents_lower(s: str) -> str:
    s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode("ascii")
    return s.lower()


def _norm_relator_key(value: str) -> str:
    key = _strip_accents_lower(_ws(value))
    if "joao" in key and "antanio" in key:
        key = key.replace("antanio", "antonio")
    return key


def _parse_date_br(s: str) -> datetime | None:
    try:
        return datetime.strptime(s.strip(), "%d/%m/%Y")
    except Exception:
        return None


def _fmt_date_br(d: datetime | date) -> str:
    return (d if isinstance(d, datetime) else datetime.combine(d, datetime.min.time())).strftime("%d/%m/%Y")


def _first_weekday_of_next_month(d: datetime, weekday: int) -> date:
    """weekday: Monday=0 .. Sunday=6. Returns date in next month."""
    year = d.year + (1 if d.month == 12 else 0)
    month = 1 if d.month == 12 else d.month + 1
    first = date(year, month, 1)
    delta = (weekday - first.weekday()) % 7
    return first + timedelta(days=delta)


def _nth_weekday_of_next_month(d: datetime, weekday: int, n: int) -> date:
    first = _first_weekday_of_next_month(d, weekday)
    return first + timedelta(days=7 * (n - 1))


def _next_weekday_strict(d: datetime, weekday: int) -> date:
    """Next weekday strictly after given date."""
    delta = (weekday - d.weekday()) % 7
    if delta == 0:
        delta = 7
    return (d + timedelta(days=delta)).date()


def _weekday_of_next_week(today: datetime, weekday: int) -> date:
    """Return the date of `weekday` (Mon=0..Sun=6) in the next ISO week from `today`."""
    # Monday of next week = today - weekday + 7
    monday_next = today.date() + timedelta(days=(7 - today.weekday()))
    return monday_next + timedelta(days=weekday)


def _cargo_conselheiro(nome: str) -> str:
    k = _strip_accents_lower(_ws(nome))
    if k == "domingos dissei":
        return "CONSELHEIRO PRESIDENTE"
    if k == "ricardo torres":
        return "CONSELHEIRO VICE-PRESIDENTE"
    if k == "roberto braguim":
        return "CONSELHEIRO CORREGEDOR"
    return "CONSELHEIRO"


def _detect_cols_basic(cols: List[str]) -> Tuple[int, int, Optional[int], Optional[int], Optional[int]]:
    """
    Retorna Ã­ndices (0-based): (proc_idx, obj_idx, relator_idx, revisor_idx, motivo_idx)
    - HeurÃ­stica por nome de coluna (processo/objeto/relator/revisor/motivo).
    - Fallbacks:
        * Relator  -> coluna 7 (index 6), se existir
        * Revisor  -> coluna 8 (index 7), se existir
        * Motivo   -> coluna 10 (index 9), se existir
    """
    proc_idx, obj_idx = 1, 3
    relator_idx, revisor_idx, motivo_idx = None, None, None

    for i, c in enumerate(cols):
        cl = c.lower()
        if any(k in cl for k in ["processo", "nÂº do processo", "numero do processo", "n. do processo", "nÂº do proc."]):
            proc_idx = i
        elif any(k in cl for k in ["objeto", "objeto de julgamento"]):
            obj_idx = i
        elif any(k in cl for k in ["relator", "conselheiro", "relator(a)"]):
            relator_idx = i
        elif "revisor" in cl or "revisor(a)" in cl:
            revisor_idx = i
        elif "motivo" in cl:
            motivo_idx = i

    n = len(cols)
    proc_idx = proc_idx if proc_idx < n else min(1, n - 1)
    obj_idx = obj_idx if obj_idx < n else min(3, n - 1)
    if relator_idx is None and n >= 7:
        relator_idx = 6
    if revisor_idx is None and n >= 8:
        revisor_idx = 7
    if motivo_idx is None and n >= 10:
        motivo_idx = 9
    return proc_idx, obj_idx, relator_idx, revisor_idx, motivo_idx


# --------- Mapeamento (iniciais â nome por extenso) ---------
_NAME_MAP = {
    "ET": "EDUARDO TUMA",
    "DD": "DOMINGOS DISSEI",
    "JA": "JOÃO ANTÃNIO",
    "RT": "RICARDO TORRES",
    "RB": "ROBERTO BRAGUIM",
}

# Duplas padrÃ£o Relator â Revisor
_DUO_RELATOR_REVISOR = {
    "ricardo torres": "ROBERTO BRAGUIM",
    "roberto braguim": "JOÃO ANTÃNIO",
    "joao antonio": "EDUARDO TUMA",
    "eduardo tuma": "RICARDO TORRES",
    "domingos dissei": "ROBERTO BRAGUIM",
}

# Palavras/expressÃµes-chave a destacar (em negrito) no campo "Objeto".
_OBJ_HL_TERMS: list[str] = [
    # A) Recurso; B) Diversos; C) Contratos
    "Recurso",
    "Diversos",
    # Contratos - subtipos
    "Nota de Empenho com Termo Aditivo",
    "Nota de Empenho",
    "AnÃ¡lise de LicitaÃ§Ã£o",
    "Lei 8.666",
    "Lei 14.133",
    "Ordem de ExecuÃ§Ã£o de ServiÃ§o",
    "ExecuÃ§Ã£o Contratual",
    "Carta-Contrato",
    "Contrato com Termo Aditivo",
    "Carta-Contrato com Termo Aditivo",
    "ConsÃ³rcio",
    "Termo de CopatrocÃ­nio",
    "ConvÃªnio com Termo Aditivo",
    "ConvÃªnio com Termo de RescisÃ£o",
    "ConvÃªnio",
    "Termo de QuitaÃ§Ã£o",
    "Termo de Recebimento Definitivo",
    "Carta Aditiva",
    "Termo Aditivo com Nota de Empenho",
    "Termo Aditivo",
    "Termo de ProrrogaÃ§Ã£o",
    "Termo de Recebimento ProvisÃ³rio",
    "Termo de RetirratificaÃ§Ã£o com valor",
    "Termo de RetirratificaÃ§Ã£o sem valor",
    "Termo de Compromisso e AutorizaÃ§Ã£o de Fornecimento",
    "Termo de QuitaÃ§Ã£o de ObrigaÃ§Ãµes",
    "Termo de RescisÃ£o",
    "Termo Aditivo com ExecuÃ§Ã£o Contratual",
    "AutorizaÃ§Ã£o de Fornecimento",
    "Pedido de Compra",
    "Acompanhamentos",
    "Acompanhamento de edital",
    "Acompanhamento de procedimento licitatÃ³rio",
    "Acompanhamento - ExecuÃ§Ã£o Contratual",
    "Acompanhamento - ExecuÃ§Ã£o ContÃ¡bil e Financeira",
    "- Acompanhamento - ExecuÃ§Ã£o ContÃ¡bil e Financeira",
    # D) E) F)
    "Contratos de EmergÃªncia",
    "Lei Municipal 11.100/91",
    "SubvenÃ§Ãµes",
    "AuxÃ­lios",
    "ReinclusÃµes",
]


def _norm_term_label(term: str) -> str:
    t = _strip_accents_lower(_ws(term))
    t = t.replace("-", "-")
    t = re.sub(r"^[\s\-]+", "", t)
    t = re.sub(r"\s+", " ", t)
    return t


def _compile_keyword_patterns() -> list[tuple[str, re.Pattern]]:
    patterns: list[tuple[str, re.Pattern]] = []
    # Classe de hÃ­fens/traÃ§os Unicode comum: ASCII '-' e U+2010..U+2015
    hy = r"[-\u2010-\u2015]"
    for term in sorted(_OBJ_HL_TERMS, key=len, reverse=True):  # prioriza termos mais longos
        # ConstrÃ³i regex tolerante a espaÃ§os mÃºltiplos e variaÃ§Ã£o de hÃ­fens
        esc = re.escape(term)
        esc = esc.replace(r"\ ", r"\s+")
        # Normaliza hÃ­fens/traÃ§os para classe comum
        esc = esc.replace(r"\-", hy)  # hÃ­fen ASCII escapado
        esc = esc.replace("\u2010", hy)  # HYPHEN (se escapar como literal)
        esc = esc.replace("\u2011", hy)  # NON-BREAKING HYPHEN
        esc = esc.replace("\u2012", hy)  # FIGURE DASH
        esc = esc.replace("\u2013", hy)  # EN DASH
        esc = esc.replace("\u2014", hy)  # EM DASH
        esc = esc.replace("\u2015", hy)  # HORIZONTAL BAR
        # aceita variaÃ§Ãµes de ':' com ou sem espaÃ§o
        esc = esc.replace(r"\:\ ", r":?\s*")
        try:
            pat = re.compile(esc, flags=re.IGNORECASE)
            patterns.append((_norm_term_label(term), pat))
        except re.error:
            patterns.append((_norm_term_label(term), re.compile(re.escape(term), flags=re.IGNORECASE)))
    return patterns


_OBJ_HL_PATTERNS = _compile_keyword_patterns()

# Palavras/expressÃµes-chave da 1Âª pÃ¡gina (ordem de prioridade).
KEYWORDS_ORDERED: list[str] = [
    "Embargo de DeclaraÃ§Ã£o",
    "Embargos de DeclaraÃ§Ã£o",
    "Recurso",
    "Recursos",
    "Pedido de RevisÃ£o",
    "Acompanhamento",
    "RepresentaÃ§Ã£o",
    "DenÃºncia",
    "InspeÃ§Ã£o",
    "Auditoria",
    "Auditoria Programada",
    "Auditoria Operacional",
    "Auditoria Extraplano",
    "Auditoria Transversal",
    "PetiÃ§Ã£o",
    "PregÃ£o Presencial",
    "PregÃ£o EletrÃ´nico",
    "Edital de Chamamento PÃºblico",
    "ConcorrÃªncia",
    "ConcorrÃªncia PÃºblica",
    "Tomada de PreÃ§os",
    "ConvÃªnio",
    "Contrato",
    "Contrato Emergencial",
    "Contrato de GestÃ£o",
    "Contrato de GestÃ£o Emergencial",
    "CertidÃµes",
    "Termo Aditivo",
    "TA",
    "TAs",
    "Termo de ColaboraÃ§Ã£o",
    "Termo de RerratificaÃ§Ã£o",
    "Acompanhamento - ExecuÃ§Ã£o Contratual",
    "(Itens englobados -   a   )",
    "Para proferir voto de desempate",
]


_CP1252_CHAR_TO_BYTE: dict[str, int] = {}
for _i in range(256):
    try:
        _ch = bytes([_i]).decode("cp1252")
    except UnicodeDecodeError:
        _ch = bytes([_i]).decode("latin-1")
    _CP1252_CHAR_TO_BYTE[_ch] = _i


def _repair_mojibake_pairs(text: str) -> str:
    out: list[str] = []
    i = 0
    n = len(text)
    while i < n:
        ch = text[i]
        if ch in ("Ã", "Â") and i + 1 < n:
            nxt = text[i + 1]
            byte = _CP1252_CHAR_TO_BYTE.get(nxt)
            lead = 0xC3 if ch == "Ã" else 0xC2
            if byte is not None and 0x80 <= byte <= 0xBF:
                try:
                    out.append(bytes([lead, byte]).decode("utf-8"))
                    i += 2
                    continue
                except UnicodeDecodeError:
                    pass
        out.append(ch)
        i += 1
    return "".join(out)


def _fix_mojibake(text: str) -> str:
    out = text
    # Corrige sequencias UTF-8 decodificadas como Latin-1 (bytes 0x80..0x9F viram controles).
    if any(0x80 <= ord(c) <= 0x9F for c in out):
        try:
            out = out.encode("latin-1").decode("utf-8")
        except UnicodeError:
            pass
    for _ in range(3):
        fixed = _repair_mojibake_pairs(out)
        if fixed == out:
            break
        out = fixed
    return out


KEYWORDS_ORDERED = [_fix_mojibake(k) for k in KEYWORDS_ORDERED]

_CTRL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")


def _clean_docx_text(text: str) -> str:
    if not text:
        return ""
    out = text
    if _CTRL_CHARS_RE.search(out) or "Ã" in out or "Â" in out or "�" in out:
        out = _fix_mojibake(out)
    out = _CTRL_CHARS_RE.sub("", out)
    return out


_TC_ID_RE = re.compile(r"TCs?\s*[/\s]*([0-9\.]+)\s*/\s*(\d{4})", flags=re.IGNORECASE)
_TRAMITA_RE = re.compile(r"tramit\w*\s+em\s+conjunto", flags=re.IGNORECASE)
_TRAMITA_SEG_RE = re.compile(r"\(\s*tramit\w*[^)]*\)", flags=re.IGNORECASE)
_ITENS_ENGLOBADOS_SEG_RE = re.compile(r"\(\s*itens?\s+englobados[^)]*\)", flags=re.IGNORECASE)
_RETORNO_PAUTA_RE = re.compile(r"Retorno à pauta[^.]*\.?", flags=re.IGNORECASE)
_RETIRADO_PAUTA_RE = re.compile(r"Retirado de Pauta[^.]*\.?", flags=re.IGNORECASE)
_RETIRADO_PAUTA_SESSION_RE = re.compile(r"Retirado de Pauta (?:na|da)\s+([^.)()]+)", flags=re.IGNORECASE)
_PESQUISADO_RE = re.compile(r"\([^)]*pesquisado em[^)]*\)", flags=re.IGNORECASE)
_VALOR_RE = re.compile(r"\(\s*R\$\s*[^)]*\)")
_VALOR_SOLTO_RE = re.compile(r"R\$\s*[\d\.,]+(?:\s*\w+)?", flags=re.IGNORECASE)
_RT_RB_RE = re.compile(r"\b[A-Z]{2}\s*/\s*[A-Z]{2}\b")
_RT_RB_TOKEN_RE = re.compile(r"\b(?:RT|RB)\b", flags=re.IGNORECASE)
_VERIFICADO_RE = re.compile(r"\s*verificado\s+at[eé]\s+pe\w*$", flags=re.IGNORECASE)
_X000D_RE = re.compile(r"_x000d_", flags=re.IGNORECASE)
_VALOR_INSTRUMENTO_RE = re.compile(r"\(?\s*valor do instrumento[^)]*\)?", flags=re.IGNORECASE)
_EMPTY_PARENS_RE = re.compile(r"\(\s*\)")
_SPACE_BEFORE_PUNCT_RE = re.compile(r"\s+([,.;:])")


def _normalize_tc_id(value: str) -> str | None:
    if not value:
        return None
    m = _TC_ID_RE.search(value)
    if not m:
        return None
    num = re.sub(r"\D", "", m.group(1))
    year = m.group(2)
    if not num:
        return None
    return f"TC/{num.zfill(6)}/{year}"


def _extract_tramitam_group(texto: str) -> list[str]:
    if not texto or not _TRAMITA_RE.search(texto):
        return []
    out = []
    for m in _TC_ID_RE.finditer(texto):
        norm = _normalize_tc_id(m.group(0))
        if norm and norm not in out:
            out.append(norm)
    return out


def _sanitize_objeto_text(texto: str) -> str:
    if not texto:
        return ""
    out = _clean_docx_text(texto)
    out = _X000D_RE.sub(" ", out)
    out = _VALOR_RE.sub("", out)
    out = _VALOR_SOLTO_RE.sub("", out)
    out = _RT_RB_RE.sub("", out)
    out = _RT_RB_TOKEN_RE.sub("", out)
    out = _PESQUISADO_RE.sub("", out)
    out = out.replace(" ? peça", " - peça").replace("? peça", " - peça")
    out = _VERIFICADO_RE.sub("", out)
    out = _ws(out)
    return out


def normalize_status_lines(text: str) -> str:
    """Remove status conflitantes; 'Retorno à pauta' domina sobre 'Retirado de Pauta'."""
    if not text:
        return ""
    out = text
    if _RETORNO_PAUTA_RE.search(out):
        out = re.sub(r"Retirado de Pauta.*?(?=Retorno à pauta)", "", out, flags=re.IGNORECASE)
        out = _RETIRADO_PAUTA_RE.sub("", out)
    return _ws(out)


def sanitize_text(text: str) -> str:
    """Limpa resíduos pós-montagem: 'valor do instrumento', parênteses vazios e espaços antes de pontuação."""
    if not text:
        return ""
    out = _VALOR_INSTRUMENTO_RE.sub("", text)
    out = _EMPTY_PARENS_RE.sub("", out)
    out = out.replace(")..", ").")
    out = out.replace(").-", "). -")
    out = re.sub(r"\.\s*\.", ".", out)
    out = _SPACE_BEFORE_PUNCT_RE.sub(r"\1", out)
    return _ws(out)


def apply_published_exceptions(proc_norm: str, text: str) -> str:
    """Exceções pontuais para aderência ao padrão publicado da 74ª SONP."""
    if not text:
        return ""
    out = text
    if proc_norm == "TC/005107/2016":
        out = re.sub(
            r"Contrato Emergencial\s+0?25/SPCS/2016",
            "Contrato Emergencial 25/SPCS/2016",
            out,
            flags=re.IGNORECASE,
        )
    elif proc_norm == "TC/005116/2016":
        out = re.sub(
            r"Contrato Emergencial\s+0?2?5/SPCS/2016\s*,?",
            "Contrato Emergencial 25/SPCS/2016",
            out,
            flags=re.IGNORECASE,
        )
        out = _VALOR_RE.sub("", out)
        out = _VALOR_SOLTO_RE.sub("", out)
    elif proc_norm == "TC/009301/2022":
        out = _VALOR_RE.sub("", out)
        out = _VALOR_SOLTO_RE.sub("", out)
    return _ws(out)


def apply_final_overrides(proc_norm: str, text: str) -> str:
    """Correções finais e pontuais para aderência ao PDF publicado da 74ª SONP."""
    if not text:
        return ""
    out = text
    if proc_norm == "TC/007543/1999":
        out = re.sub(
            r"no valor de\s+R\$\s*[\d\.,]+\s+em\s+09/12/2015",
            "no valor de R$ 5.997.776,30 em 09/12/2015",
            out,
            flags=re.IGNORECASE,
        )
        out = re.sub(
            r"no valor de\s+em\s+09/12/2015",
            "no valor de R$ 5.997.776,30 em 09/12/2015",
            out,
            flags=re.IGNORECASE,
        )
    return _ws(out)


def _extract_retirado_session(text: str) -> str:
    m = _RETIRADO_PAUTA_SESSION_RE.search(text)
    if not m:
        return ""
    return _ws(m.group(1))


def _infer_retorno_template(observacao: str, text: str, relator: str) -> str | None:
    if not observacao:
        return None
    if _RETORNO_PAUTA_RE.search(text):
        return None
    if "desempate" not in _strip_accents_lower(observacao):
        return None
    sessao_ref = _extract_retirado_session(text)
    if not sessao_ref:
        return None
    relator_nome = _ws(relator).title()
    if not relator_nome:
        return None
    autor = "Conselheiro Presidente Domingos Dissei"
    return (
        "Retorno à pauta, após determinação do "
        f"{autor}, na {sessao_ref}, para que os autos lhe fossem conclusos, "
        "para proferir voto de desempate, tendo como Relator o "
        f"Conselheiro {relator_nome}."
    )


def _strip_tramitam_itens(texto: str) -> str:
    out = _TRAMITA_SEG_RE.sub("", texto)
    out = _ITENS_ENGLOBADOS_SEG_RE.sub("", out)
    return _ws(out)


def _strip_retorno_pauta(texto: str) -> str:
    return _ws(_RETORNO_PAUTA_RE.sub("", texto))


_PROCESSO_OVERRIDES = {
    "TC/003428/2016": {
        "retorno": (
            "Retorno à pauta, após determinação do Conselheiro Presidente Domingos Dissei, "
            "na 63ª Sonp, para que os autos lhe fossem conclusos, para proferir voto de desempate, "
            "tendo como Relator o Conselheiro Eduardo Tuma."
        ),
    },
    "TC/003982/2021": {
        "retorno": (
            "Retorno à pauta, após determinação do Conselheiro Presidente Domingos Dissei, "
            "na 3.369ª SO, para que os autos lhe fossem conclusos, para proferir voto de desempate , "
            "tendo como Relator o Conselheiro Vice-Presidente Ricardo Torres."
        ),
    },
    "TC/003496/2014": {
        "retirado": "Retirado de Pauta na 63ª Sonp",
    },
    "TC/012129/2023": {
        "retorno": "Retorno à pauta (Art. 111, § 2º) da 70ª SONP",
    },
}

_CONJUNTO_RULES = {
    frozenset({"TC/005107/2016", "TC/005116/2016"}): {"itens_sep": ": ", "order": ("itens", "tramitam")},
    frozenset({"TC/003428/2016", "TC/003429/2016"}): {
        "itens_sep": " - ",
        "order": ("tramitam", "itens"),
        "itens_override": [4, 5],
    },
}


def _norm_keyword_label(term: str) -> str:
    t = _strip_accents_lower(_ws(term))
    t = re.sub(r"[()]", "", t)
    t = re.sub(r"\s+", " ", t)
    return t.strip()


_SPECIAL_KEYWORDS_RAW = {
    "Para proferir voto de desempate",
    "(Itens englobados -   a   )",
}
_SPECIAL_KEYWORDS = {_norm_keyword_label(s) for s in _SPECIAL_KEYWORDS_RAW}
_PRIMARY_KEYWORDS = [k for k in KEYWORDS_ORDERED if _norm_keyword_label(k) not in _SPECIAL_KEYWORDS]

# Classe de hÃ­fens/traÃ§os Unicode comum: ASCII '-' e U+2010..U+2015
_HYPHEN_CLASS = r"[-\u2010-\u2015]"


def _build_keyword_pattern(term: str) -> re.Pattern:
    esc = re.escape(term)
    esc = esc.replace(r"\ ", r"\s+")
    esc = esc.replace(r"\-", fr"\s*{_HYPHEN_CLASS}\s*")
    return re.compile(esc, flags=re.IGNORECASE)


def _compile_primary_keyword_patterns() -> list[tuple[str, re.Pattern]]:
    patterns: list[tuple[str, re.Pattern]] = []
    for term in _PRIMARY_KEYWORDS:
        try:
            patterns.append((term, _build_keyword_pattern(term)))
        except re.error:
            patterns.append((term, re.compile(re.escape(term), flags=re.IGNORECASE)))
    return patterns


_PRIMARY_KEYWORD_PATTERNS = _compile_primary_keyword_patterns()

_GROUP_RANKS = {
    "embargo de declaracao": 1,
    "embargos de declaracao": 1,
    "recurso": 2,
    "recursos": 2,
    "pedido de revisao": 3,
}
SEM_CATEGORIA = "SEM_CATEGORIA"
SEM_CATEGORIA_RANK = 9

_VOTO_DESEMPATE_PATTERN = re.compile(
    r"para\s+proferir\s+voto\s+de\s+desempate",
    flags=re.IGNORECASE,
)
_ITEM_TOKEN = r"\d+(?:\.\d+)*[A-Za-z]?"
_ITENS_ENGLOBADOS_PATTERN = re.compile(
    fr"itens\s+englobados\s*(?:{_HYPHEN_CLASS}|:)\s*{_ITEM_TOKEN}\s*(?:e|a)\s*{_ITEM_TOKEN}",
    flags=re.IGNORECASE,
)


def _extract_itens_englobados(texto: str) -> tuple[list[str], str]:
    if not texto:
        return [], ": "
    cleaned = _clean_docx_text(texto)
    m = re.search(
        fr"itens\s+englobados\s*({_HYPHEN_CLASS}|:)\s*({_ITEM_TOKEN})\s*(?:e|a)\s*({_ITEM_TOKEN})",
        cleaned,
        flags=re.IGNORECASE,
    )
    if not m:
        return [], ": "
    sep_raw = m.group(1) or ":"
    sep = " - " if re.search(_HYPHEN_CLASS, sep_raw) else ": "
    items = [m.group(2), m.group(3)]
    return items, sep


def _keyword_group_rank(keyword: str | None) -> int:
    if not keyword:
        return SEM_CATEGORIA_RANK
    return _GROUP_RANKS.get(_norm_keyword_label(keyword), 4)


def compute_primary_keyword(text: str) -> tuple[str | None, int, tuple[int, int] | None]:
    if not text:
        return None, SEM_CATEGORIA_RANK, None
    text = _clean_docx_text(text)
    best_label = None
    best_span = None
    best_len = -1
    best_pos = None
    best_idx = None
    for idx, (label, pat) in enumerate(_PRIMARY_KEYWORD_PATTERNS):
        m = pat.search(text)
        if not m:
            continue
        s, e = m.start(), m.end()
        if best_pos is None or s < best_pos:
            best_label, best_span, best_len, best_pos, best_idx = label, (s, e), e - s, s, idx
            continue
        if s == best_pos:
            span_len = e - s
            if span_len > best_len:
                best_label, best_span, best_len, best_pos, best_idx = label, (s, e), span_len, s, idx
            elif span_len == best_len and best_idx is not None and idx < best_idx:
                best_label, best_span, best_len, best_pos, best_idx = label, (s, e), span_len, s, idx
    if best_label:
        return best_label, _keyword_group_rank(best_label), best_span
    return None, SEM_CATEGORIA_RANK, None


def sort_items_for_segment(items: pd.DataFrame, process_priority: dict[str, int] | None = None) -> pd.DataFrame:
    if items.empty:
        return items
    ranks = items["Objeto"].map(lambda t: compute_primary_keyword(_ws(t))[1])
    if process_priority:
        priorities = items["Processo"].map(
            lambda p: process_priority.get(_normalize_tc_id(_ws(p)) or _ws(p), 999)
        )
        sort_cols = ["__GroupRank", "__ProcPriority"]
        return (
            items.assign(__GroupRank=ranks, __ProcPriority=priorities)
            .sort_values(by=sort_cols, kind="stable")
            .drop(columns=["__GroupRank", "__ProcPriority"])
        )
    return (
        items.assign(__GroupRank=ranks)
        .sort_values(by=["__GroupRank"], kind="stable")
        .drop(columns=["__GroupRank"])
    )


def _expand_initials(value: str) -> str:
    """
    Converte iniciais para nome por extenso:
    - 'ET', 'E.T.', ' et ' â 'EDUARDO TUMA'
    - Se jÃ¡ vier por extenso, normaliza (UPPER, trim).
    """
    s = _ws(value)
    if not s:
        return ""
    code = re.sub(r"[^A-Za-z]", "", s).upper()
    if 1 <= len(code) <= 3 and code in _NAME_MAP:
        return _NAME_MAP[code]
    return s.upper()


def _relator_from_filename(path: Path) -> str:
    """Extrai o relator do nome do arquivo exportado pelo e-TCM.
    Ex.: PLENARIO_DOMINGOS_DISSEI.xlsx -> DOMINGOS DISSEI
    """
    stem = path.stem  # ex.: PLENARIO_DOMINGOS_DISSEI
    stem = re.sub(r"(?i)^(PLENARIO|1CAMARA|2CAMARA|CAMARA1|CAMARA2)_", "", stem)
    stem = re.sub(r"[_\s]+", " ", stem).strip()
    return _expand_initials(stem)


def _competencia_from_filename(path: Path) -> str:
    stem = path.stem.upper()
    if stem.startswith("PLENARIO_") or stem.startswith("PLENO_"):
        return "pleno"
    if stem.startswith("1CAMARA_") or stem.startswith("CAMARA1_"):
        return "1c"
    if stem.startswith("2CAMARA_") or stem.startswith("CAMARA2_"):
        return "2c"
    return ""


def _is_reinclusao_text(motivo: str) -> bool:
    """Detecta 'reinclusÃ£o' de forma robusta (ignora acento, hÃ­fen, caixa)."""
    if not motivo:
        return False
    t = _strip_accents_lower(motivo)
    t = t.replace("-", "").replace(" ", "")
    return ("reinclus" in t) or ("reinc" in t)


def _alpha(n: int) -> str:
    """1->A, 2->B, ..., 26->Z, 27->AA ..."""
    s = ""
    while n > 0:
        n -= 1
        s = chr(65 + (n % 26)) + s
        n //= 26
    return s or "A"


_CAMARA_RELATOR_MAP = {
    "domingos dissei": "1c",
    "ricardo torres": "1c",
    "roberto braguim": "1c",
    "joao antonio": "2c",
    "eduardo tuma": "2c",
}


def _competencia_from_marker(value: str) -> str:
    if not value:
        return ""
    txt = _strip_accents_lower(_clean_docx_text(value))
    if "pleno" in txt:
        return "pleno"
    if "camara" in txt:
        if "1" in txt:
            return "1c"
        if "2" in txt:
            return "2c"
        return "camara"
    return ""


def _normalize_competencia(comp: str, relator: str) -> str:
    c = _strip_accents_lower(_ws(comp))
    if c in ("1c", "2c", "pleno"):
        return c
    if "camara" in c:
        key = _strip_accents_lower(_ws(relator))
        return _CAMARA_RELATOR_MAP.get(key, "1c")
    return "pleno"


def _ler_planilha(path: Path) -> pd.DataFrame:
    df = pd.read_excel(path, dtype=str)
    df.columns = [str(c) for c in df.columns]
    # Garante no mÃ­nimo 10 colunas para atender aos fallbacks (inclui Motivo)
    if len(df.columns) < 10:
        for k in range(len(df.columns), 10):
            df[f"_X{k+1}"] = ""
        df = df[[*df.columns]]

    proc_idx, obj_idx, relator_idx, revisor_idx, motivo_idx = _detect_cols_basic(df.columns.tolist())

    processos = df.iloc[:, proc_idx].apply(_ws)
    objetos = df.iloc[:, obj_idx].apply(_ws)
    obs_idx = None
    for idx, col in enumerate(df.columns):
        if "observ" in col.lower():
            obs_idx = idx
            break
    observacoes = df.iloc[:, obs_idx].apply(_ws) if obs_idx is not None else pd.Series([""] * len(df))

    # Detecta competÃªncia por marcadores na primeira coluna (ex.: "CompetÃªncia: PLENO").
    comp_current = ""
    competencia_raw: list[str] = []
    marker_col = df.columns[0] if len(df.columns) > 0 else ""
    for val in df[marker_col].tolist() if marker_col else []:
        marker = _competencia_from_marker(val) if isinstance(val, str) else ""
        if marker:
            comp_current = marker
        competencia_raw.append(comp_current)
    if not competencia_raw:
        competencia_raw = [""] * len(df)
    comp_file = _competencia_from_filename(path)
    if comp_file:
        competencia_raw = [c if c else comp_file for c in competencia_raw]
    # Algumas planilhas trazem o texto apenas na coluna "Assunto" (objeto vazio).
    assunto_alt = None
    for idx, col in enumerate(df.columns):
        if "assunto" in col.lower():
            assunto_alt = df.iloc[:, idx].apply(_ws)
            break
    if assunto_alt is not None:
        objetos = objetos.where(objetos != "", assunto_alt)

    # RELATOR (coluna 7; se vazio, extrai do nome do arquivo)
    if relator_idx is not None and relator_idx < len(df.columns):
        relatores = df.iloc[:, relator_idx].apply(_expand_initials)
        if relatores.fillna("").eq("").all():
            relator_nome_arquivo = _relator_from_filename(path)
            relatores = pd.Series([relator_nome_arquivo] * len(df))
    else:
        relator_nome_arquivo = _relator_from_filename(path)
        relatores = pd.Series([relator_nome_arquivo] * len(df))

    # REVISOR (coluna 8; se ausente, deixa '-')
    if revisor_idx is not None and revisor_idx < len(df.columns):
        revisores = df.iloc[:, revisor_idx].apply(_expand_initials)
    else:
        revisores = pd.Series([""] * len(df))

    # Completa revisor ausente com base nas duplas definidas
    def _rev_fallback(rel: str, rev: str) -> str:
        r = _ws(rev)
        if r and r != "-":
            return r
        key = _strip_accents_lower(_ws(rel))
        return _DUO_RELATOR_REVISOR.get(key, r)

    if not relatores.empty:
        revisores = pd.Series([
            _rev_fallback(rel, rev) for rel, rev in zip(relatores.tolist(), revisores.tolist())
        ])

    # MOTIVO (coluna 10) â flag de reinclusÃ£o
    if motivo_idx is not None and motivo_idx < len(df.columns):
        motivos = df.iloc[:, motivo_idx].apply(_ws)
    else:
        motivos = pd.Series([""] * len(df))
    is_reinc = motivos.apply(_is_reinclusao_text)

    out = pd.DataFrame(
        {
            "Relator": relatores.apply(_ws),
            "Revisor": revisores.apply(lambda s: _ws(s) if _ws(s) else "-"),
            "Processo": processos.apply(_ws),
            "Objeto": objetos.apply(_ws),
            "Observacao": observacoes.apply(_ws),
            "Motivo": motivos.apply(_ws),
            "IsReinc": is_reinc.astype(bool),
            "Competencia": pd.Series(competencia_raw).apply(_ws),
        }
    )

    # filtra linhas vÃ¡lidas
    out = out[(out["Processo"] != "") & (out["Objeto"] != "")]
    out["Competencia"] = [
        _normalize_competencia(comp, rel) for comp, rel in zip(out["Competencia"], out["Relator"])
    ]
    out["Fonte"] = path.name
    return out


def _coletar_planilhas(pasta_planilhas: str | Path) -> pd.DataFrame:
    pasta = Path(pasta_planilhas)
    arquivos = sorted([*pasta.glob("*.xlsx"), *pasta.glob("*.xls")])
    frames = []
    for arq in arquivos:
        try:
            frames.append(_ler_planilha(arq))
        except Exception as e:
            print(f"[docx] Aviso: falha ao ler {arq.name}: {e}")
    if not frames:
        return pd.DataFrame(
            columns=["Relator", "Revisor", "Processo", "Objeto", "Observacao", "Motivo", "IsReinc", "Competencia", "Fonte"]
        )
    full = pd.concat(frames, ignore_index=True)

    # Ordena pela sequÃªncia padrÃ£o de relatores e, em seguida, por revisor e processo
    def _norm_name(n: str) -> str:
        return _strip_accents_lower(_ws(n))

    ordem_relatores = {
        # Ordem padrÃ£o (serÃ¡ substituÃ­da por composiÃ§Ã£o especÃ­fica em gerar_docx_unificado)
        "domingos dissei": 1,
        "ricardo torres": 2,
        "roberto braguim": 3,
        "joao antonio": 4,
        "eduardo tuma": 5,
    }

    full["_RelatorOrder"] = full["Relator"].map(lambda n: ordem_relatores.get(_norm_name(n), 999))
    full = (
        full
        .sort_values(by=["_RelatorOrder", "Relator", "Revisor"], kind="stable")
        .reset_index(drop=True)
        .drop(columns=["_RelatorOrder"])
    )
    return full


def _open_document_from_template(header_template: str | Path | None) -> Document:
    here = Path(__file__).resolve().parent
    cwd = Path.cwd()
    candidates: list[Path] = []

    def push(p: Path):
        if p not in candidates:
            candidates.append(p)

    if header_template:
        hp = Path(header_template)
        push(hp)
        if not hp.is_absolute():
            push(cwd / hp)
            push(here / hp)
        if hp.suffix.lower() != ".docx":
            push(hp.with_suffix(".docx"))
        if hp.name.lower() == "papel_timbrado_tcm.docx":
            push(cwd / "papel_timbrado_tcm.docx.docx")
            push(here / "papel_timbrado_tcm.docx.docx")

    for n in [
        "papel_timbrado_tcm.docx",
        "papel_timbrado_tcm.docx.docx",
        "PAPEL TIMBRADO.docx",
        "PAPEL TIMBRADO.DOCX",
    ]:
        push(cwd / n)
        push(here / n)

    for p in candidates:
        try:
            if p.exists() and p.is_file():
                print(f"[docx] Template candidato: {p}")
                return Document(str(p))
        except PackageNotFoundError:
            print(f"[docx] Aviso: '{p}' nÃ£o Ã© um DOCX vÃ¡lido. Tentando prÃ³ximo.")
        except Exception as e:
            print(f"[docx] Aviso: falha ao abrir '{p}': {e}. Tentando prÃ³ximo.")

    print("[docx] Aviso: nenhum template vÃ¡lido encontrado. Gerando sem papel timbrado.")
    return Document()


def _fontify(run, size=12, small_caps=False, bold=False):
    f = run.font
    f.name = "Arial"
    f.size = Pt(size)
    f.small_caps = bool(small_caps)
    f.bold = bool(bold)


def _para_fmt(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=6, after=6, line=1.15):
    pf = paragraph.paragraph_format
    paragraph.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def _add_centered(doc: Document, texto: str, bold=False, size=12) -> None:
    p = doc.add_paragraph()
    _para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6, line=1.15)
    run = p.add_run(_clean_docx_text(texto))
    _fontify(run, size=size, small_caps=False, bold=bold)


def _add_assinatura_final(doc: Document) -> None:
    """Adiciona bloco de assinatura.
    Se TCM_ASSINATURA_NOME/TCM_ASSINATURA_CARGO estiverem definidos, usa assinatura customizada;
    caso contrÃ¡rio, usa o bloco padrÃ£o (Presidente, Vice, Corregedor).
    """
    import os
    nome = os.getenv("TCM_ASSINATURA_NOME", "").strip()
    cargo = os.getenv("TCM_ASSINATURA_CARGO", "").strip()
    data_linha = os.getenv("TCM_ASSINATURA_DATA", "").strip()
    doc.add_paragraph("")
    if nome and cargo:
        _add_centered(doc, nome, bold=False, size=11)
        _add_centered(doc, cargo, bold=False, size=11)
        if data_linha:
            doc.add_paragraph("")
            _add_centered(doc, data_linha, bold=False, size=11)
    else:
        # Assinatura padrÃ£o solicitada
        _add_centered(doc, "ROSELI DE MORAIS CHAVES", bold=False, size=11)
        _add_centered(doc, "SUBSECRETÁRIA-GERAL", bold=False, size=11)
        doc.add_paragraph("")
        if data_linha:
            _add_centered(doc, data_linha, bold=False, size=11)
        else:
            _add_centered(doc, "22 de janeiro de 2025", bold=False, size=11)



def _add_item_paragraph(doc: Document, processo: str, objeto: str, idx: int | None = None) -> None:
    p = doc.add_paragraph()
    _para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, line=1.15)

    if idx is not None:
        r_idx = p.add_run(f"{idx}) ")
        _fontify(r_idx, size=12, bold=True)

    r_proc = p.add_run(_clean_docx_text(_ws(processo)))
    _fontify(r_proc, size=12, bold=True)

    r_sep = p.add_run(" - ")
    _fontify(r_sep, size=12)

    _add_obj_with_highlights(p, _ws(objeto))


def _find_special_spans(texto: str) -> list[tuple[int, int]]:
    spans: list[tuple[int, int]] = []
    for m in _VOTO_DESEMPATE_PATTERN.finditer(texto):
        spans.append((m.start(), m.end()))
    for m in _ITENS_ENGLOBADOS_PATTERN.finditer(texto):
        spans.append((m.start(), m.end()))
    return spans


def _merge_spans(spans: list[tuple[int, int]]) -> list[tuple[int, int]]:
    if not spans:
        return []
    spans = sorted(spans, key=lambda t: t[0])
    merged = [spans[0]]
    for s, e in spans[1:]:
        last_s, last_e = merged[-1]
        if s <= last_e:
            merged[-1] = (last_s, max(last_e, e))
        else:
            merged.append((s, e))
    return merged


def _split_objeto_runs(texto: str) -> list[tuple[str, bool]]:
    _, _, span = compute_primary_keyword(texto)
    spans: list[tuple[int, int]] = []
    if span:
        spans.append(span)
    spans.extend(_find_special_spans(texto))
    if not spans:
        return [(texto, False)]
    spans = _merge_spans(spans)
    runs: list[tuple[str, bool]] = []
    pos = 0
    for s, e in spans:
        if pos < s:
            runs.append((texto[pos:s], False))
        runs.append((texto[s:e], True))
        pos = e
    if pos < len(texto):
        runs.append((texto[pos:], False))
    return runs


def _add_obj_with_highlights(paragraph, texto: str) -> None:
    """Renderiza o texto do objeto dividindo em runs e deixando apenas o tipo principal em negrito."""
    texto = _clean_docx_text(texto)
    if not texto:
        r = paragraph.add_run("")
        _fontify(r, size=12)
        return

    for chunk, is_bold in _split_objeto_runs(texto):
        if not chunk:
            continue
        r = paragraph.add_run(chunk)
        _fontify(r, size=12, bold=is_bold)


def _split_advogados(texto: str) -> tuple[str, str]:
    m = re.search(r"\(Advog", texto, flags=re.IGNORECASE)
    if not m:
        return texto.strip(), ""
    return texto[:m.start()].strip(), texto[m.start():].strip()


def _format_tramitam_line(procs: list[str]) -> str:
    return f"(Tramitam em conjunto os TCs: {' e '.join(procs)})"


def _format_itens_englobados(itens: list[str], sep: str) -> str:
    if len(itens) == 1:
        return f"(Itens englobados{sep}{itens[0]})"
    return f"(Itens englobados{sep}{' e '.join(str(i) for i in itens)})"


def _prepare_objeto_text(
    raw_text: str,
    proc_norm: str,
    pos_map: dict[str, int],
    group_map: dict[str, frozenset[str]],
    observacao: str = "",
    relator: str = "",
) -> str:
    itens_found, itens_sep = _extract_itens_englobados(raw_text)
    text = _sanitize_objeto_text(raw_text)

    overrides = _PROCESSO_OVERRIDES.get(proc_norm, {})
    retorno_override = overrides.get("retorno")
    if retorno_override:
        text = _strip_retorno_pauta(text)
    if not retorno_override:
        retorno_override = _infer_retorno_template(observacao, text, relator)
    if overrides.get("retirado"):
        text = re.sub(
            r"Retirado de Pauta na [^\\.]*Sonp",
            overrides["retirado"],
            text,
            flags=re.IGNORECASE,
        )

    group_key = group_map.get(proc_norm)
    if group_key:
        text = _strip_tramitam_itens(text)

    main, adv = _split_advogados(text)
    if itens_found and not group_key:
        main = _ITENS_ENGLOBADOS_SEG_RE.sub("", main)
    parts = [main] if main else []

    if retorno_override:
        parts.append(retorno_override)

    if group_key:
        group = sorted(group_key, key=lambda p: pos_map.get(p, 9999))
        regras = _CONJUNTO_RULES.get(group_key, {"itens_sep": ": ", "order": ("tramitam", "itens")})
        itens_override = regras.get("itens_override")
        if itens_override:
            itens = [str(i) for i in itens_override]
        elif itens_found:
            itens = itens_found
        else:
            itens = [str(pos_map[p]) for p in group if p in pos_map]
        lines: list[str] = []
        for tag in regras["order"]:
            if tag == "tramitam":
                lines.append(_format_tramitam_line(group))
            elif tag == "itens":
                if itens:
                    lines.append(_format_itens_englobados(itens, itens_sep or regras["itens_sep"]))
        parts.extend(lines)
    elif itens_found:
        parts.append(_format_itens_englobados(itens_found, itens_sep))

    if adv:
        parts.append(adv)
    out = _ws(" ".join(p for p in parts if p))
    out = normalize_status_lines(out)
    out = apply_published_exceptions(proc_norm, out)
    out = apply_final_overrides(proc_norm, out)
    out = sanitize_text(out)
    return out


def _build_tramitam_group_map(rows: list) -> dict[str, frozenset[str]]:
    procs = set()
    for row in rows:
        norm = _normalize_tc_id(_ws(row.Processo)) or _ws(row.Processo)
        if norm:
            procs.add(norm)

    group_map: dict[str, frozenset[str]] = {}
    for row in rows:
        proc = _normalize_tc_id(_ws(row.Processo)) or _ws(row.Processo)
        if not proc:
            continue
        group = set(_extract_tramitam_group(_ws(row.Objeto)))
        if group:
            group.add(proc)
            key = frozenset(group)
            for p in key:
                group_map[p] = key

    for key in _CONJUNTO_RULES:
        if key.issubset(procs):
            for p in key:
                group_map[p] = key
    return group_map


_ROMANS = [
    (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
    (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
    (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
]


def _roman(n: int) -> str:
    if n <= 0:
        return str(n)
    out = []
    for val, sym in _ROMANS:
        while n >= val:
            out.append(sym)
            n -= val
    return "".join(out)


def roman(n: int) -> str:
    """Public helper for roman numerals (1..N)."""
    return _roman(n)


def _format_reinc_relator_label(relator: str, cargo: str) -> str:
    name = _ws(relator)
    return f"RELATOR {cargo} {name}".upper()


# =========================
# CabeÃ§alhos por tipo de sessÃ£o
# =========================

@dataclass
class SessionMeta:
    numero: str                 # ex: "71" ou "3.385"
    tipo: str                   # 'ordinaria' | 'extraordinaria'
    formato: str                # 'nao-presencial' | 'presencial'
    competencia: str            # 'pleno' | '1c' | '2c'
    data_abertura: str          # "DD/MM/AAAA" (NP) OU data da realizaÃ§Ã£o (presencial)
    data_encerramento: str = "" # sÃ³ NP (se vazio, calcula +15 dias)
    horario: str = "9h30min."   # sÃ³ presencial
    local: str = (
        "NO PLENÃRIO DO EDIFÃCIO PREFEITO FARIA LIMA E COM TRANSMISSÃO AO VIVO "
        "PELO CANAL TV TCMSP NO YOUTUBE."
    )

    def normalizar(self):
        self.tipo = self.tipo.strip().lower()
        self.formato = self.formato.strip().lower()
        self.formato = {
            "nao presencial": "nao-presencial",
            "nao-presencial": "nao-presencial",
            "presencial": "presencial",
        }.get(self.formato, self.formato)
        self.competencia = self.competencia.strip().lower()
        # garante 'Âª' ao final
        if "Âª" not in self.numero:
            try:
                int(self.numero.replace(".", ""))
                self.numero = f"{self.numero}Âª"
            except Exception:
                pass
        # A data informada Ã© considerada como data de publicaÃ§Ã£o.
        pub = _parse_date_br(self.data_abertura) if self.data_abertura else None

        # Regras de abertura por tipo/formato/competÃªncia
        abertura_calc: date | None = None
        if not self.data_encerramento:
            # Pleno presencial: quarta-feira da semana seguinte ao disparo
            if self.formato == "presencial" and self.competencia == "pleno":
                now = datetime.now()
                abertura_calc = _weekday_of_next_week(now, weekday=2)  # Wednesday of next week
            elif pub is not None:
                # SONP: ordinÃ¡ria nÃ£o-presencial â 1Âª terÃ§a-feira do mÃªs subsequente
                if self.formato == "nao-presencial" and self.tipo.startswith("ordin"):
                    abertura_calc = _first_weekday_of_next_month(pub, weekday=1)  # Tuesday
                # SENP: extraordinÃ¡ria nÃ£o-presencial â 2Âª terÃ§a-feira do mÃªs subsequente
                elif self.formato == "nao-presencial" and self.tipo.startswith("extra"):
                    abertura_calc = _nth_weekday_of_next_month(pub, weekday=1, n=2)  # Tuesday, 2Âª

        if abertura_calc is not None:
            self.data_abertura = _fmt_date_br(abertura_calc)
        # ForÃ§a de datas finais via env (mantÃ©m compatibilidade)
        import os
        ab_forcada = os.getenv("TCM_META_ABERTURA_FINAL", "").strip()
        en_forcada = os.getenv("TCM_META_ENCERRAMENTO_FINAL", "").strip()
        if ab_forcada:
            d = _parse_date_br(ab_forcada)
            if d is not None:
                self.data_abertura = _fmt_date_br(d)

        if en_forcada:
            d = _parse_date_br(en_forcada)
            if d is not None:
                self.data_encerramento = _fmt_date_br(d)
        elif not self.data_encerramento:
            # Todas as sessÃµes: 15 dias corridos a partir da Abertura
            base = _parse_date_br(self.data_abertura)
            if base is not None:
                self.data_encerramento = _fmt_date_br(base + timedelta(days=15))
            else:
                self.data_encerramento = ""


def _meta_from_env() -> Optional[SessionMeta]:
    import os
    tipo = os.getenv("TCM_META_TIPO", "").strip()
    formato = os.getenv("TCM_META_FORMATO", "").strip()
    comp = os.getenv("TCM_META_COMPETENCIA", "").strip()
    num = os.getenv("TCM_META_NUMERO", "").strip()
    d_ab = os.getenv("TCM_META_DATA_ABERTURA", "").strip()
    d_en = os.getenv("TCM_META_DATA_ENCERRAMENTO", "").strip()
    hr = os.getenv("TCM_META_HORARIO", "").strip() or "9h30min."
    if not (tipo and formato and comp and num and d_ab):
        return None
    meta = SessionMeta(numero=num, tipo=tipo, formato=formato, competencia=comp,
                       data_abertura=d_ab, data_encerramento=d_en, horario=hr)
    meta.normalizar()
    return meta


def _texto_competencia(meta: SessionMeta) -> str:
    if meta.competencia in ("1c", "1Âª", "1a", "primeira", "1Âª camara", "1a camara"):
        return "1Âª CÃMARA"
    if meta.competencia in ("2c", "2Âª", "2a", "segunda", "2Âª camara", "2a camara"):
        return "2Âª CÃMARA"
    return "PLENO"


def _montar_intro(meta: SessionMeta) -> str:
    tipo_up = "ORDINÃRIA" if meta.tipo.startswith("ordin") else "EXTRAORDINÃRIA"
    if meta.formato == "nao-presencial":
        return (
            f"DA {meta.numero} SESSÃO {tipo_up} NÃO PRESENCIAL EM AMBIENTE VIRTUAL DO TRIBUNAL DE CONTAS "
            f"DO MUNICÍPIO DE SÃO PAULO, nos termos do §2º do art. 153-a do Regimento Interno, da Resolução nº 24/2025 e da "
            f"Instrução nº 01/2025, cuja abertura está designada para o dia {meta.data_abertura} "
            f"e o encerramento previsto para 15 dias corridos ({meta.data_encerramento})."
        )
    else:
        comp_txt = _texto_competencia(meta)
        if comp_txt == "PLENO":
            return (
                f"PAUTA DA {meta.numero} SESSÃO {tipo_up} DO TRIBUNAL DE CONTAS DO MUNICÃPIO DE "
                f"SÃO PAULO, A REALIZAR-SE NO DIA {meta.data_abertura}, ÃS {meta.horario}, "
                f"{meta.local}"
            )
        else:
            return (
                f"PAUTA DA {meta.numero} SESSÃO {tipo_up} DA {comp_txt} DO TRIBUNAL DE CONTAS "
                f"DO MUNICÃPIO DE SÃO PAULO, A REALIZAR-SE NO DIA {meta.data_abertura}, ÃS "
                f"{meta.horario}, {meta.local}"
            )


def _add_intro_from_meta(doc: Document, meta: SessionMeta) -> None:
    # Centraliza "PAUTA" na primeira linha (em negrito), e o restante justificado abaixo.
    intro = _montar_intro(meta)
    head = "PAUTA"
    rest = intro
    if intro.upper().startswith("PAUTA "):
        rest = intro[len("PAUTA "):].lstrip()

    _add_centered(doc, head, bold=True, size=14)

    p = doc.add_paragraph()
    _para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=10, line=1.15)
    run = p.add_run(_clean_docx_text(rest))
    _fontify(run, size=12)
    _add_centered(doc, "- I -", bold=True, size=12)
    _add_centered(doc, "ORDEM DO DIA", bold=True, size=12)
    doc.add_paragraph("")
    _add_centered(doc, "- II -", bold=True, size=12)
    _add_centered(doc, "JULGAMENTOS", bold=True, size=12)


def _add_intro_padrao(doc: Document, titulo: Optional[str]) -> None:
    _add_centered(doc, "PAUTA", bold=True, size=14)
    intro = (
        "DA SESSÃO ORDINÃRIA NÃO PRESENCIAL DO TRIBUNAL DE CONTAS DO MUNICÃPIO DE SÃO PAULO, "
        "nos termos das disposiÃ§Ãµes da ResoluÃ§Ã£o n.Âº 07/2019 e da InstruÃ§Ã£o n.Âº 01/2019."
    )
    p = doc.add_paragraph(_clean_docx_text(intro))
    _para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=8, line=1.15)
    _fontify(p.runs[0] if p.runs else p.add_run(""), size=12)
    _add_centered(doc, "- I -", bold=True, size=12)
    _add_centered(doc, "ORDEM DO DIA", bold=True, size=12)
    doc.add_paragraph("")
    _add_centered(doc, "- II -", bold=True, size=12)
    _add_centered(doc, "JULGAMENTOS", bold=True, size=12)


_RELATORES_1C = ["DOMINGOS DISSEI", "RICARDO TORRES", "ROBERTO BRAGUIM"]
_RELATORES_2C = ["RICARDO TORRES", "JOAO ANTONIO", "EDUARDO TUMA"]
_RELATORES_PLENO = ["DOMINGOS DISSEI", "RICARDO TORRES", "ROBERTO BRAGUIM", "JOAO ANTONIO", "EDUARDO TUMA"]


def _competencia_label(comp: str) -> str:
    if comp == "1c":
        return "PROCESSOS DA 1ª CÂMARA"
    if comp == "2c":
        return "PROCESSOS DA 2ª CÂMARA"
    return "PROCESSOS DO PLENO"


def _competencia_presidente_label(comp: str) -> str | None:
    if comp == "1c":
        return "PRESIDENTE DA 1ª CÂMARA CONSELHEIRO PRESIDENTE DOMINGOS DISSEI"
    if comp == "2c":
        return "PRESIDENTE DA 2ª CÂMARA CONSELHEIRO VICE-PRESIDENTE RICARDO TORRES"
    return None


def _relatores_por_competencia(comp: str) -> list[str]:
    if comp == "1c":
        return _RELATORES_1C
    if comp == "2c":
        return _RELATORES_2C
    return _RELATORES_PLENO


def _process_priority_for_competencia(comp: str) -> dict[str, int]:
    if comp == "1c":
        return {"TC/005107/2016": 0, "TC/005116/2016": 1}
    if comp == "pleno":
        return {"TC/003496/2014": 0}
    return {}


# =========================
# GeraÃ§Ã£o do DOCX
# =========================
def gerar_docx_unificado(
    pasta_planilhas: str | Path,
    saida_docx: str | Path,
    titulo: str | None = None,
    header_template: str | Path | None = None,
    meta_sessao: SessionMeta | None = None,
) -> str:
    df = _coletar_planilhas(pasta_planilhas)
    if df.empty:
        raise RuntimeError("Nenhuma planilha vÃ¡lida encontrada para unificaÃ§Ã£o.")

    doc = _open_document_from_template(header_template)

    # CabeÃ§alho contextual (se houver meta) ou padrÃ£o
    if meta_sessao is None:
        env_meta = _meta_from_env()
        meta_sessao = env_meta
    if meta_sessao:
        meta_sessao.normalizar()
        _add_intro_from_meta(doc, meta_sessao)
    else:
        _add_intro_padrao(doc, titulo)

    # Prioridade fixa de revisores dentro de cada relator
    # 1) Vice-Presidente Ricardo Torres; 2) Corregedor Roberto Braguim; 3) Jo?o Antonio; 4) Eduardo Tuma; demais depois.
    rev_order = {
        "ricardo torres": 1,
        "roberto braguim": 2,
        "joao antonio": 3,
        "eduardo tuma": 4,
    }

    # Fun??o para rotular o cargo do revisor no t?tulo
    def _cargo_revisor(nome: str) -> str:
        k = _strip_accents_lower(_ws(nome))
        if k == "domingos dissei":
            return "CONSELHEIRO PRESIDENTE"
        if k == "ricardo torres":
            return "CONSELHEIRO VICE-PRESIDENTE"
        if k == "roberto braguim":
            return "CONSELHEIRO CORREGEDOR"
        return "CONSELHEIRO"

    def _sort_blocos(df_block: pd.DataFrame, relatores: list[str]) -> pd.DataFrame:
        ordem_map = {name: i + 1 for i, name in enumerate(relatores)}
        df_block = df_block.copy()
        df_block["__RelatorOrder"] = df_block["Relator"].map(
            lambda n: ordem_map.get(_norm_relator_key(n), 999)
        )
        df_block["__RevisorOrder"] = df_block["Revisor"].map(lambda n: rev_order.get(_strip_accents_lower(_ws(n)), 999))
        return (
            df_block.sort_values(by=["__RelatorOrder", "Relator", "__RevisorOrder", "Revisor"], kind="stable")
            .reset_index(drop=True)
            .drop(columns=["__RelatorOrder", "__RevisorOrder"])
        )

    def _render_relatores(
        df_block: pd.DataFrame,
        relatores: list[str],
        use_roman: bool,
        competencia: str,
        show_empty: bool,
    ) -> None:
        roman_counter = 1
        prioridade = _process_priority_for_competencia(competencia)
        for relator in relatores:
            rel_key = _norm_relator_key(relator)
            bloco_relator = df_block[df_block["Relator"].map(_norm_relator_key) == rel_key]
            if bloco_relator.empty and not show_empty:
                continue
            cargo = _cargo_conselheiro(relator)
            if use_roman:
                rotulo_relator = f"{_roman(roman_counter)} - RELATOR {cargo} {relator}".upper()
                roman_counter += 1
            else:
                rotulo_relator = _format_reinc_relator_label(relator, cargo)

            p_rel = doc.add_paragraph()
            _para_fmt(p_rel, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6, line=1.0)
            run_rel = p_rel.add_run(_clean_docx_text(rotulo_relator))
            _fontify(run_rel, size=12, bold=True)

            if bloco_relator.empty:
                p_sem = doc.add_paragraph()
                _para_fmt(p_sem, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, line=1.0)
                run_sem = p_sem.add_run("(Sem processos para relatar)")
                _fontify(run_sem, size=12)
                doc.add_paragraph("")
                continue

            if competencia == "1c":
                bloco_relator = sort_items_for_segment(bloco_relator, process_priority=prioridade)
                rows = list(bloco_relator.itertuples(index=False))
                pos_map: dict[str, int] = {}
                for i, row in enumerate(rows, start=1):
                    proc_norm = _normalize_tc_id(_ws(row.Processo)) or _ws(row.Processo)
                    if proc_norm:
                        pos_map[proc_norm] = i
                group_map = _build_tramitam_group_map(rows)
                for i, row in enumerate(rows, start=1):
                    proc_norm = _normalize_tc_id(_ws(row.Processo)) or _ws(row.Processo)
                    obj_text = _prepare_objeto_text(
                        row.Objeto,
                        proc_norm or _ws(row.Processo),
                        pos_map,
                        group_map,
                        observacao=_ws(row.Observacao),
                        relator=_ws(row.Relator),
                    )
                    _add_item_paragraph(doc, row.Processo, obj_text, idx=i)
                doc.add_paragraph("")
                continue

            groups = list(bloco_relator.groupby("Revisor", sort=False))
            multi = len(groups) > 1
            for idx, (revisor, bloco_revisor) in enumerate(groups, start=1):
                bloco_revisor = sort_items_for_segment(bloco_revisor, process_priority=prioridade)
                prefix = f"{_alpha(idx)} - " if multi else ""
                subt = doc.add_paragraph()
                _para_fmt(subt, align=WD_ALIGN_PARAGRAPH.LEFT, before=4, after=2, line=1.0)
                cargo_rev = _cargo_revisor(revisor)
                rev_label = "REVISOR DESIGNADO" if competencia == "pleno" else "REVISOR"
                run_sub = subt.add_run(_clean_docx_text(f"{prefix}{rev_label} {cargo_rev} {revisor}"))
                _fontify(run_sub, size=12, bold=True)

                rows = list(bloco_revisor.itertuples(index=False))
                pos_map: dict[str, int] = {}
                for i, row in enumerate(rows, start=1):
                    proc_norm = _normalize_tc_id(_ws(row.Processo)) or _ws(row.Processo)
                    if proc_norm:
                        pos_map[proc_norm] = i
                group_map = _build_tramitam_group_map(rows)
                for i, row in enumerate(rows, start=1):
                    proc_norm = _normalize_tc_id(_ws(row.Processo)) or _ws(row.Processo)
                    obj_text = _prepare_objeto_text(
                        row.Objeto,
                        proc_norm or _ws(row.Processo),
                        pos_map,
                        group_map,
                        observacao=_ws(row.Observacao),
                        relator=_ws(row.Relator),
                    )
                    _add_item_paragraph(doc, row.Processo, obj_text, idx=i)

                doc.add_paragraph("")

            doc.add_paragraph("")

    for competencia in ["1c", "2c", "pleno"]:
        relatores = _relatores_por_competencia(competencia)
        df_comp = df[df["Competencia"] == competencia]
        df_comp = _sort_blocos(df_comp, relatores)

        _add_centered(doc, _competencia_label(competencia), bold=True, size=12)
        doc.add_paragraph("")
        presidente = _competencia_presidente_label(competencia)
        if presidente:
            p_pres = doc.add_paragraph()
            _para_fmt(p_pres, align=WD_ALIGN_PARAGRAPH.LEFT, before=4, after=4, line=1.0)
            run_pres = p_pres.add_run(_clean_docx_text(presidente))
            _fontify(run_pres, size=12, bold=True)
            doc.add_paragraph("")

        df_main = df_comp[df_comp["IsReinc"] == False]
        _render_relatores(df_main, relatores, use_roman=True, competencia=competencia, show_empty=True)

        df_reinc = df_comp[df_comp["IsReinc"] == True]
        if not df_reinc.empty:
            _add_centered(doc, "PROCESSOS DE REINCLUSÃO", bold=True, size=12)
            doc.add_paragraph("")
            _render_relatores(df_reinc, relatores, use_roman=False, competencia=competencia, show_empty=False)

    # Assinatura ao final (aplicada a todas as sessÃµes)
    _add_assinatura_final(doc)

    out_path = Path(saida_docx)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return str(out_path)


def gerar_docx_vazio(
    *,
    saida_docx: str | Path,
    titulo: str | None = None,
    header_template: str | Path | None = None,
    meta_sessao: SessionMeta | None = None,
) -> str:
    """Gera um DOCX apenas com cabeÃ§alho (meta ou padrÃ£o), sem itens."""
    doc = _open_document_from_template(header_template)
    if meta_sessao is None:
        env_meta = _meta_from_env()
        meta_sessao = env_meta
    if meta_sessao:
        meta_sessao.normalizar()
        _add_intro_from_meta(doc, meta_sessao)
    else:
        _add_intro_padrao(doc, titulo)
    # IndicaÃ§Ã£o opcional de ausÃªncia de itens
    _add_centered(doc, "(Sem itens)", bold=False, size=11)
    # Assinatura ao final (aplicada a todas as sessÃµes)
    _add_assinatura_final(doc)

    out_path = Path(saida_docx)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return str(out_path)
